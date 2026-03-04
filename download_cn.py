# 1. 医疗 - 来源: HuggingFace Mirror "https://huggingface.co/datasets/Aunderline/CMeEE-V2/resolve/main/CMeEE-V2_train.js# 新代码 (GitHub 源，更稳定且无需 Token)
# 2. 新闻 - 来源: OYE93 Chinese-NLP-Corpus (BIO格式） "https://raw.githubusercontent.com/OYE93/Chinese-NLP-Corpus/master/NER/MSRA/msra_train_bio.txt"
# 3. 简历 - 来源：https://github.com/jiesutd/LatticeLSTM/tree/master/ResumeNER/train.char.bmes
# 4. 生活 - 来源：https://github.com/hltcoe/golden-horse/tree/master/data/weiboNER_2nd_conll.train


import os
import docx


WORK_DIR = "./input_processcn"
FILE_WEIBO = "weiboNER_2nd_conll.train"  # 生活化数据
FILE_RESUME = "train.char.bmes"  # 简历数据


def save_to_word(texts, filename_prefix, title):
    """将提取到的文本保存为 Word"""
    if not texts:
        print(f" {filename_prefix} 数据为空，跳过。")
        return
    texts = list(set([t.strip() for t in texts if len(t.strip()) > 2]))
    total = len(texts)
    print(f"提取到 {total} 条有效数据，正在生成 Word...")

    # 分卷保存
    BATCH_SIZE = 3000
    for i in range(0, total, BATCH_SIZE):
        batch_texts = texts[i: i + BATCH_SIZE]
        doc = docx.Document()
        doc.add_heading(f'{title} (Part {i // BATCH_SIZE + 1})', 0)

        for t in batch_texts:
            # 清洗非法字符
            clean_text = "".join([c for c in t if c.isprintable() or c in ['\n', '\t', '\r']])
            doc.add_paragraph(clean_text)

        # 保存到 WORK_DIR (即 input_processcn 目录)
        filename = f"{filename_prefix}_{i // BATCH_SIZE + 1:02d}.docx"
        save_path = os.path.join(WORK_DIR, filename)
        doc.save(save_path)
        print(f"已保存: {os.path.abspath(save_path)}")


def process_bio_file(filename, output_prefix, title):
    """通用解析器：适用于字+标签格式 (BIO/BMES)"""
    file_path = os.path.join(WORK_DIR, filename)
    print(f"[解析] 正在读取: {filename} ...")

    if not os.path.exists(file_path):
        print(f" 错误: 找不到文件 {file_path}")
        return

    texts = []
    current_chars = []

    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if not line:
                    if current_chars:
                        sentence = "".join(current_chars)
                        if len(sentence) > 1:
                            texts.append(sentence)
                        current_chars = []
                else:
                    parts = line.split()
                    if parts:
                        current_chars.append(parts[0])

            if current_chars:
                texts.append("".join(current_chars))

        save_to_word(texts, output_prefix, title)

    except Exception as e:
        print(f" 解析出错: {e}")



def main():
    print(f">>>开始处理文件夹: {os.path.abspath(WORK_DIR)}")

    if not os.path.exists(WORK_DIR):
        print(f" 错误: 找不到目录 {WORK_DIR}")
        return
    process_bio_file(FILE_RESUME, "Dataset_HR_Resume", "中文简历人事数据")

    print("---------------------------------------------------------------")

    process_bio_file(FILE_WEIBO, "Dataset_Life_Weibo", "中文生活化微博数据")

    print("\n>>> 全部完成！请打开 input_processcn 文件夹查看生成的 Word 文档。")


if __name__ == "__main__":
    main()