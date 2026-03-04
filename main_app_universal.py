# -*- coding: utf-8 -*-
"""
@Description:
    多模态文档脱敏核心引擎。
    基于局部可预测性(Predictability)与结构稳定性(Stability)的评估，
    结合上下文角色(Role)感知模型，实现 L1-L5 多粒度的隐私脱敏与加密映射策略。
"""

import os
import re
import torch
import sqlite3
import hashlib
import docx
import openpyxl
import math
import time
import fitz  # PyMuPDF
from transformers import AutoTokenizer, AutoModelForTokenClassification
from Crypto.Cipher import AES

# ================= Configuration =================
NER_MODEL_PATH = "model_ner_multilingual"
DB_PATH = "universal_mapping.db"
INPUT_FOLDER = "./input_file"
OUTPUT_FOLDER = "./output_file"
MASTER_KEY = b"12345678901234567890123456789012"

# === Algorithm Hyperparameters ===
TAU_P = 3.0  # Predictability 阈值
TAU_S = 0.8  # Stability 阈值
K_ANONYMITY = 3  # K-匿名严格度系数
ALPHA = 0.1  # Laplace 平滑参数 (用于语料统计)


class SecureDB:
    """
    基于 SQLite 的本地安全映射库。
    使用 AES-256 (EAX 模式) 对原始明文进行加密存储，保障数据不可逆转性。
    """

    def __init__(self, db_path, key):
        self.db_path = db_path
        self.key = key
        self.conn = sqlite3.connect(self.db_path, check_same_thread=False)
        self.cursor = self.conn.cursor()
        self.init_db()
        self.pending_commits = 0

    def init_db(self):
        self.cursor.execute('PRAGMA journal_mode=WAL;')
        self.cursor.execute('PRAGMA synchronous=OFF;')
        self.cursor.execute('''CREATE TABLE IF NOT EXISTS forward_index (
                        hash_key TEXT PRIMARY KEY, 
                        proxy_value TEXT,
                        entity_type TEXT)''')
        self.cursor.execute('''CREATE TABLE IF NOT EXISTS secure_storage (
                        token TEXT PRIMARY KEY, 
                        encrypted_real_value BLOB, 
                        nonce BLOB)''')
        self.cursor.execute('''CREATE TABLE IF NOT EXISTS corpus_stats (
                        gram TEXT PRIMARY KEY, 
                        count INTEGER)''')
        self.cursor.execute('''CREATE TABLE IF NOT EXISTS structure_stats (
                        entity_type TEXT,
                        position INTEGER,
                        char_type TEXT,
                        count INTEGER,
                        PRIMARY KEY (entity_type, position, char_type))''')
        self.cursor.execute('''CREATE TABLE IF NOT EXISTS qid_history (
                        qid_hash TEXT PRIMARY KEY,
                        frequency INTEGER)''')
        self.conn.commit()

    def close(self):
        if self.conn:
            self.conn.commit()
            self.conn.close()

    def manual_commit(self):
        self.conn.commit()
        self.pending_commits = 0

    def save_mapping(self, real_value, proxy_value, entity_type, role="DEFAULT"):
        """将实体文本与生成的代理值进行加密绑定并存入数据库"""
        unique_str = f"{real_value}_{role}"
        token = hashlib.sha256(unique_str.encode()).hexdigest()[:16]
        hash_key = hashlib.sha256(unique_str.encode()).hexdigest()

        cipher = AES.new(self.key, AES.MODE_EAX)
        ciphertext, tag = cipher.encrypt_and_digest(real_value.encode('utf-8'))

        self.cursor.execute("INSERT OR IGNORE INTO forward_index VALUES (?, ?, ?)",
                            (hash_key, proxy_value, entity_type))
        self.cursor.execute("INSERT OR REPLACE INTO secure_storage VALUES (?, ?, ?)",
                            (token, ciphertext + tag, cipher.nonce))
        self.pending_commits += 1
        return token

    def update_corpus_count(self, text):
        if not text: return
        self.cursor.execute(
            "INSERT INTO corpus_stats (gram, count) VALUES (?, 1) ON CONFLICT(gram) DO UPDATE SET count=count+1",
            (text,))
        self.cursor.execute(
            "INSERT INTO corpus_stats (gram, count) VALUES ('__TOTAL__', 1) ON CONFLICT(gram) DO UPDATE SET count=count+1")

    def get_corpus_count(self, text):
        c = self.cursor.execute("SELECT count FROM corpus_stats WHERE gram=?", (text,)).fetchone()
        t = self.cursor.execute("SELECT count FROM corpus_stats WHERE gram='__TOTAL__'").fetchone()
        return (c[0] if c else 0), (t[0] if t else 1)

    def update_structure_stats(self, entity_type, signature):
        for idx, char_type in enumerate(signature):
            self.cursor.execute(
                "INSERT INTO structure_stats (entity_type, position, char_type, count) VALUES (?, ?, ?, 1) ON CONFLICT(entity_type, position, char_type) DO UPDATE SET count=count+1",
                (entity_type, idx, char_type))

    def get_structure_probs(self, entity_type, length):
        data = {}
        cursor = self.cursor.execute("SELECT position, char_type, count FROM structure_stats WHERE entity_type=?",
                                     (entity_type,))
        for pos, ctype, cnt in cursor:
            if pos >= length: continue
            if pos not in data: data[pos] = {}
            data[pos][ctype] = cnt
        return data

    def check_and_update_qid(self, proxy_val):
        qid_hash = hashlib.md5(proxy_val.encode()).hexdigest()
        res = self.cursor.execute("SELECT frequency FROM qid_history WHERE qid_hash=?", (qid_hash,)).fetchone()
        freq = res[0] if res else 0
        self.cursor.execute(
            "INSERT INTO qid_history (qid_hash, frequency) VALUES (?, 1) ON CONFLICT(qid_hash) DO UPDATE SET frequency=frequency+1",
            (qid_hash,))
        return freq


class ContextAnalyzer:
    """上下文角色分析模块：基于周围文本特征分配 Role，缓解同名实体的语义冲突"""

    def __init__(self):
        self.role_rules = {
            'PER': {
                'DOC': ['医生', '医师', '大夫', '主任', '专家', '教授', '医护'],
                'PAT': ['患者', '病人', '病患', '确诊', '住院', '家属'],
                'TEACHER': ['老师', '教师', '导师', '班主任'],
                'STUDENT': ['学生', '同学', '班长']
            },
            'ORG': {
                'HOSPITAL': ['医院', '诊所', '卫生院'],
                'COMPANY': ['公司', '集团', '企业'],
                'SCHOOL': ['大学', '中学', '小学', '学院']
            }
        }

    def infer_role(self, text_segment, entity_text, entity_type, window_size=20):
        if entity_type not in self.role_rules: return "DEFAULT"
        try:
            start_idx = text_segment.find(entity_text)
            if start_idx == -1: return "DEFAULT"

            left = max(0, start_idx - window_size)
            right = min(len(text_segment), start_idx + len(entity_text) + window_size)
            context = text_segment[left:right]

            rules = self.role_rules[entity_type]
            for role, keywords in rules.items():
                for kw in keywords:
                    if kw in context: return role
            return "DEFAULT"
        except:
            return "DEFAULT"


class SemanticAnalyzer:
    """语义特征提取模块：计算 P值(局部可预测性) 与 S值(结构稳定性)"""

    def __init__(self, db):
        self.db = db

    def calculate_predictability(self, text):
        count_cx, total_c = self.db.get_corpus_count(text)
        V = 10000
        numerator = count_cx + ALPHA
        denominator = total_c + ALPHA * V
        prob = numerator / denominator
        return -math.log(prob) if prob > 0 else 0.0

    def generate_signature(self, text):
        sig = []
        for char in text:
            if char.isdigit():
                sig.append('D')
            elif char.isalpha():
                sig.append('L')
            else:
                sig.append('S')
        return sig

    def calculate_stability(self, text, entity_type):
        signature = self.generate_signature(text)
        length = len(signature)
        if length == 0: return 0.0

        stats = self.db.get_structure_probs(entity_type, length)
        self.db.update_structure_stats(entity_type, signature)

        if not stats: return 1.0
        match_score = 0
        for i, char_type in enumerate(signature):
            if i in stats:
                counts = stats[i]
                most_common_type = max(counts, key=counts.get)
                if char_type == most_common_type: match_score += 1
        return match_score / length


class PrivacyDecisionEngine:
    """隐私决策引擎：根据实体特征进行多级代理值(Proxy)生成"""

    def __init__(self, ner_path, mask_model, db):
        self.db = db
        self.analyzer = SemanticAnalyzer(db)
        self.context_analyzer = ContextAnalyzer()
        self.device = "cuda" if torch.cuda.is_available() else "cpu"
        print(f"[INFO] Loading NLP model on device: {self.device}")

        self.tokenizer = AutoTokenizer.from_pretrained(ner_path)
        self.model = AutoModelForTokenClassification.from_pretrained(ner_path).to(self.device)
        self.id2label = self.model.config.id2label

        self.vocab = {
            'SURNAMES': ["王", "李", "张", "刘", "陈", "杨", "黄", "赵", "周", "吴"],
            'NAMES': ["伟", "芳", "娜", "敏", "静", "丽", "强", "磊", "军", "洋"],
            'CITIES': ["北京市", "上海市", "广州市", "深圳市", "杭州市"],
            'DISTRICTS': ["朝阳区", "海淀区", "浦东区", "天河区", "余杭区"]
        }
        self.l5_map = {
            'PER': '匿名用户', 'LOC': '某地', 'ORG': '某机构',
            'PHONE': '联系方式', 'EMAIL': '通信端点', 'ID_CARD': '身份凭证'
        }

    def generate_proxy_final(self, text, entity_type, role="DEFAULT"):
        if not text: return text
        p_score = self.analyzer.calculate_predictability(text)
        s_score = self.analyzer.calculate_stability(text, entity_type)
        self.db.update_corpus_count(text)

        # 四象限动态映射规则
        level = 5
        if p_score >= TAU_P:
            level = 2 if s_score >= TAU_S else 4
        else:
            level = 3 if s_score >= TAU_S else 5

        # K-匿名约束检查与动态升维
        current_proxy = ""
        while level <= 5:
            current_proxy = self._generate_strategy(level, text, entity_type, role)
            freq = self.db.check_and_update_qid(current_proxy)
            if freq >= K_ANONYMITY:
                break
            else:
                if level == 5: break
                level += 1
        return current_proxy

    def _generate_strategy(self, level, text, entity_type, role):
        if level == 1:
            return text

        elif level == 2:  # L2: 局部遮蔽 (Masking)
            length = len(text)
            if length <= 2: return text[0] + "*"
            if entity_type == 'PHONE': return text[:3] + "****" + text[-4:]
            if entity_type == 'ID_CARD': return text[:6] + "********" + text[-4:]
            if entity_type == 'EMAIL':
                parts = text.split('@')
                if len(parts) == 2: return parts[0][:2] + "***@" + parts[1]
            return text[:2] + "*" * (length - 4) + text[-2:] if length > 4 else text[0] + "**"

        elif level == 3:  # L3: 结构泛化 (Structural Generalization)
            h = hashlib.md5((text + role).encode()).hexdigest()[:4]
            return f"[{entity_type}#{h}]"

        elif level == 4:  # L4: 类别拟真替换 (Category Replacement)
            seed_str = text + role
            h_int = int(hashlib.md5(seed_str.encode()).hexdigest(), 16)
            suffix = hashlib.md5(seed_str.encode()).hexdigest()[:3]

            if entity_type == 'PER':
                s = self.vocab['SURNAMES'][h_int % len(self.vocab['SURNAMES'])]
                n = self.vocab['NAMES'][(h_int // 10) % len(self.vocab['NAMES'])]
                return f"{s}{n}_{suffix}"
            elif entity_type == 'LOC':
                c = self.vocab['CITIES'][h_int % len(self.vocab['CITIES'])]
                return f"{c}_{suffix}"

            return f"[{entity_type}_{len(text)}]"

        else:  # L5: 完全抽象 (Complete Abstraction)
            tag = self.l5_map.get(entity_type, entity_type)
            # 引入微型指纹 (Micro-Fingerprint)，解决多对一映射导致的不可逆问题
            short_h = hashlib.md5((text + role).encode()).hexdigest()[:2]
            return f"[{tag}_{short_h}]"


class DocumentProcessor:
    """多模态文档解析与替换处理器"""

    def __init__(self, ner_path, mask_path, db):
        self.db = db
        self.engine = PrivacyDecisionEngine(ner_path, mask_path, db)

        self.patterns = [
            (re.compile(r'1[3-9]\d{9}'), 'PHONE'),
            (re.compile(r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+'), 'EMAIL'),
            (re.compile(r'\d{17}[\dX]'), 'ID_CARD')
        ]

        # 停用词及领域高频词白名单，防止模型过拟合误报
        self.blocklist = {
            'Mon', 'Tue', 'Wed', 'Thu', 'Fri', 'Sat', 'Sun',
            'Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec',
            'Date', 'Subject', 'To', 'From', 'Sent', 'Cc', 'Bcc', 'Forwarded', 'Original', 'Message',
            'Enron', 'Email', 'Dataset', 'Batch'
        }

    def process_batch_texts(self, text_list):
        if not text_list: return []

        final_texts = []
        for text in text_list:
            if not text or not text.strip():
                final_texts.append(text)
                continue

            temp_text = text

            # 1. 优先执行正则表达式匹配
            replacements = []
            for pat, p_type in self.patterns:
                for match in pat.finditer(temp_text):
                    val = match.group()
                    role = self.engine.context_analyzer.infer_role(temp_text, val, p_type)
                    proxy = self.engine.generate_proxy_final(val, p_type, role)
                    self.db.save_mapping(val, proxy, p_type, role)
                    replacements.append((match.start(), match.end(), proxy))

            replacements.sort(key=lambda x: x[0], reverse=True)
            for start, end, proxy in replacements:
                temp_text = temp_text[:start] + proxy + temp_text[end:]

            # 2. 执行模型级命名实体识别 (NER)
            inputs = self.engine.tokenizer([temp_text], return_tensors="pt", truncation=True, max_length=512).to(
                self.engine.device)
            with torch.no_grad():
                outputs = self.engine.model(**inputs)
            preds = torch.argmax(outputs.logits, dim=2)[0]
            tokens = self.engine.tokenizer.convert_ids_to_tokens(inputs["input_ids"][0])

            curr_w, curr_t = "", None
            ner_ents = set()
            for tok, idx in zip(tokens, preds):
                lbl = self.engine.id2label[idx.item()]
                cln = tok.replace("##", "")
                if tok in ['[CLS]', '[SEP]', '[PAD]']: continue
                if lbl.startswith("B-"):
                    if curr_w: ner_ents.add((curr_w, curr_t))
                    curr_w, curr_t = cln, lbl.split("-")[1]
                elif lbl.startswith("I-") and curr_t:
                    curr_w += cln
            if curr_w: ner_ents.add((curr_w, curr_t))

            for w, t in ner_ents:
                if len(w) < 2 or w in self.blocklist: continue
                if w not in temp_text: continue
                if w.isdigit(): continue

                role = self.engine.context_analyzer.infer_role(temp_text, w, t)
                proxy = self.engine.generate_proxy_final(w, t, role)
                self.db.save_mapping(w, proxy, t, role)

                # TODO: 当前采用简单的字符替换，未来需优化为基于 Offset 的精准替换以避免潜在冲突
                temp_text = temp_text.replace(w, proxy, 1)

            final_texts.append(temp_text)

        return final_texts

    def process_text_content(self, text):
        return self.process_batch_texts([text])[0]

    def process_docx(self, input_path, output_path):
        try:
            print(f"[INFO] Parsing DOCX file: {os.path.basename(input_path)} ...")
            doc = docx.Document(input_path)
            BATCH_SIZE = 32
            tasks = []

            for i, para in enumerate(doc.paragraphs):
                if para.text.strip(): tasks.append({'index': i, 'text': para.text})

            total_tasks = len(tasks)
            start_time = time.time()

            for i in range(0, total_tasks, BATCH_SIZE):
                batch_tasks = tasks[i: i + BATCH_SIZE]
                texts = [t['text'] for t in batch_tasks]
                processed_texts = self.process_batch_texts(texts)

                for j, task in enumerate(batch_tasks):
                    doc.paragraphs[task['index']].text = processed_texts[j]

                self.db.manual_commit()

                if (i + BATCH_SIZE) % 160 == 0 or (i + BATCH_SIZE) >= total_tasks:
                    current = min(i + BATCH_SIZE, total_tasks)
                    elapsed = time.time() - start_time
                    if elapsed > 0:
                        speed = current / elapsed
                        remain = (total_tasks - current) / speed
                        print(
                            f"       Progress: {current}/{total_tasks} | Speed: {speed:.1f} para/s | ETA: {remain / 60:.1f} min")

            if len(doc.tables) > 0:
                print("       Processing embedded tables...")
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                cell.text = self.process_text_content(cell.text)
                self.db.manual_commit()

            out_dir = os.path.dirname(output_path)
            if not os.path.exists(out_dir): os.makedirs(out_dir)
            doc.save(output_path)
            print(f"[SUCCESS] Exported: {os.path.basename(output_path)}")

        except Exception as e:
            print(f"[ERROR] DOCX processing failed: {e}")

    def process_excel(self, input_path, output_path):
        try:
            print(f"[INFO] Parsing EXCEL file: {os.path.basename(input_path)} ...")
            wb = openpyxl.load_workbook(input_path)
            BATCH_SIZE = 64
            tasks = []

            for ws in wb:
                for row in ws.iter_rows():
                    for cell in row:
                        if isinstance(cell.value, str) and cell.value.strip():
                            tasks.append({'cell': cell, 'text': cell.value})

            total = len(tasks)
            start_time = time.time()

            for i in range(0, total, BATCH_SIZE):
                batch = tasks[i: i + BATCH_SIZE]
                texts = [t['text'] for t in batch]
                results = self.process_batch_texts(texts)
                for j, task in enumerate(batch):
                    task['cell'].value = results[j]

                self.db.manual_commit()

                if (i + BATCH_SIZE) % 2000 == 0:
                    curr = min(i + BATCH_SIZE, total)
                    sp = curr / (time.time() - start_time)
                    print(f"       Progress: {curr}/{total} | Speed: {sp:.1f} cells/s")

            out_dir = os.path.dirname(output_path)
            if not os.path.exists(out_dir): os.makedirs(out_dir)
            wb.save(output_path)
            print(f"[SUCCESS] Exported: {os.path.basename(output_path)}")
        except Exception as e:
            print(f"[ERROR] EXCEL processing failed: {e}")

    def process_pdf(self, input_path, output_path):
        try:
            print(f"[INFO] Extracting PDF content: {os.path.basename(input_path)} ...")
            doc = fitz.open(input_path)
            new_doc = docx.Document()
            BATCH_SIZE = 16
            pages_text = []

            for page in doc:
                t = page.get_text()
                if t.strip(): pages_text.append(t)

            processed_pages = []
            for i in range(0, len(pages_text), BATCH_SIZE):
                batch_texts = pages_text[i: i + BATCH_SIZE]
                results = self.process_batch_texts(batch_texts)
                processed_pages.extend(results)
                self.db.manual_commit()
                print(f"       Progress: {min(i + BATCH_SIZE, len(pages_text))}/{len(pages_text)} pages")

            for txt in processed_pages: new_doc.add_paragraph(txt)
            output_docx = output_path + ".docx"
            out_dir = os.path.dirname(output_docx)
            if not os.path.exists(out_dir): os.makedirs(out_dir)
            new_doc.save(output_docx)
            print(f"[SUCCESS] Reconstructed as DOCX: {os.path.basename(output_docx)}")
        except Exception as e:
            print(f"[ERROR] PDF processing failed: {e}")


def main():
    if not os.path.exists(INPUT_FOLDER): os.makedirs(INPUT_FOLDER)
    if not os.path.exists(OUTPUT_FOLDER): os.makedirs(OUTPUT_FOLDER)

    print("[SYSTEM] Initialization started...")
    db = SecureDB(DB_PATH, MASTER_KEY)
    processor = DocumentProcessor(NER_MODEL_PATH, "unused", db)

    files = [f for f in os.listdir(INPUT_FOLDER) if not f.startswith("~$")]
    print(f"[SYSTEM] Detected {len(files)} target files.")

    for filename in files:
        in_path = os.path.join(INPUT_FOLDER, filename)
        out_path = os.path.join(OUTPUT_FOLDER, "脱敏_" + filename)
        ext = filename.lower().split('.')[-1]

        if ext in ['docx', 'doc']:
            processor.process_docx(in_path, out_path)
        elif ext in ['xlsx', 'xls']:
            processor.process_excel(in_path, out_path)
        elif ext == 'pdf':
            processor.process_pdf(in_path, out_path)
        else:
            print(f"[WARNING] Skipping unsupported format: {filename}")

    db.close()
    print("[SYSTEM] Execution completed.")


if __name__ == "__main__":
    main()