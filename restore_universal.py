# -*- coding: utf-8 -*-
"""
@Description: 
    多模态文档语义还原引擎。
    通过读取本地映射数据库 (SQLite)，利用预编译正则表达式在段落级别进行精确匹配。
    结合 AES-256 解密机制完成信息的安全还原，并自动记录审计日志。
"""

import os
import re
import sqlite3
import docx
import openpyxl
import hashlib
import time
from Crypto.Cipher import AES

# ================= Configuration =================
DB_PATH = "universal_mapping.db"
INPUT_FOLDER = "./output_file"
OUTPUT_FOLDER = "./output_file"
MASTER_KEY = b"12345678901234567890123456789012"


class Restorer:
    """脱敏数据还原控制类"""

    def __init__(self, db_path, key):
        self.db_path = db_path
        self.key = key

        # 建立数据库持久连接，支持 GUI 异步调用
        self.conn = sqlite3.connect(self.db_path, check_same_thread=False)
        self.cursor = self.conn.cursor()

        self.init_audit_table()
        self.mapping_cache = self.load_all_mappings()

        # 动态编译高并发正则匹配引擎
        if self.mapping_cache:
            print("[INFO] Compiling regex engine for restoration...")
            # 按长度倒序排序，避免短模式截断长模式 (如 Person_1 截断 Person_12)
            sorted_keys = sorted(self.mapping_cache.keys(), key=len, reverse=True)
            safe_keys = [re.escape(k) for k in sorted_keys]
            pattern_str = '|'.join(safe_keys)
            self.regex_pattern = re.compile(pattern_str)
            print(f"[INFO] Engine initialized. Entity patterns loaded: {len(safe_keys)}")
        else:
            self.regex_pattern = None

    def init_audit_table(self):
        """建立系统级审计表，记录追踪还原操作行为"""
        self.cursor.execute('''CREATE TABLE IF NOT EXISTS restoration_log (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        filename TEXT,
                        restore_time TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                        status TEXT,
                        note TEXT)''')
        self.conn.commit()

    def close(self):
        if self.conn:
            self.conn.close()

    def log_operation(self, filename, status, note=""):
        try:
            self.cursor.execute("INSERT INTO restoration_log (filename, status, note) VALUES (?, ?, ?)",
                                (filename, status, note))
            self.conn.commit()
        except Exception as e:
            print(f"[WARNING] Audit log failed: {e}")

    def decrypt(self, ciphertext, nonce, tag):
        try:
            cipher = AES.new(self.key, AES.MODE_EAX, nonce=nonce)
            return cipher.decrypt_and_verify(ciphertext, tag).decode('utf-8')
        except:
            return None

    def load_all_mappings(self):
        """将正向索引加载至内存，降低 I/O 损耗"""
        if not os.path.exists(self.db_path):
            print(f"[ERROR] Database not found: {self.db_path}")
            return {}

        try:
            self.cursor.execute("SELECT hash_key, proxy_value FROM forward_index")
            rows = self.cursor.fetchall()
        except:
            return {}

        mapping = {}
        collision_count = 0

        for row in rows:
            token = row[0][:16]
            proxy = row[1]
            if proxy in mapping:
                collision_count += 1
            mapping[proxy] = token

        if collision_count > 0:
            print(f"[WARNING] Detected {collision_count} potential mapping collisions (due to L5 Abstraction).")

        return mapping

    def restore_value(self, proxy_text):
        if proxy_text not in self.mapping_cache: return None
        token = self.mapping_cache[proxy_text]

        row = self.cursor.execute("SELECT encrypted_real_value, nonce FROM secure_storage WHERE token=?",
                                  (token,)).fetchone()

        if row:
            return self.decrypt(row[0][:-16], row[1], row[0][-16:])
        return None

    def _replace_func(self, match):
        proxy = match.group(0)
        real = self.restore_value(proxy)
        # 解密失败则视为潜在篡改，保持原状输出
        return real if real else proxy

    def process_text_content(self, text):
        if not text or not self.regex_pattern:
            return text
        return self.regex_pattern.sub(self._replace_func, text)

    def get_output_filename(self, filename):
        if filename.startswith("脱敏_"): return filename.replace("脱敏_", "还原_", 1)
        return "还原_" + filename

    def process_docx(self, input_path, output_path):
        try:
            print(f"[INFO] Restoring DOCX: {os.path.basename(input_path)} ...")
            doc = docx.Document(input_path)
            tasks = []

            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    tasks.append({'index': i, 'text': para.text})

            total = len(tasks)
            BATCH_SIZE = 100
            start_time = time.time()

            for i in range(0, total, BATCH_SIZE):
                batch = tasks[i: i + BATCH_SIZE]
                for task in batch:
                    restored = self.process_text_content(task['text'])
                    doc.paragraphs[task['index']].text = restored

                if (i + BATCH_SIZE) % 500 == 0:
                    curr = min(i + BATCH_SIZE, total)
                    elapsed = time.time() - start_time
                    speed = curr / elapsed if elapsed > 0 else 0
                    print(f"       Progress: {curr}/{total} | Speed: {speed:.1f} para/s")

            if len(doc.tables) > 0:
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                cell.text = self.process_text_content(cell.text)

            out_dir = os.path.dirname(output_path)
            if not os.path.exists(out_dir): os.makedirs(out_dir)
            doc.save(output_path)
            print(f"[SUCCESS] Saved: {os.path.basename(output_path)}")
            self.log_operation(os.path.basename(input_path), "SUCCESS", "DOCX Restoration successful")
            return True

        except Exception as e:
            print(f"[ERROR] DOCX restoration failed: {e}")
            self.log_operation(os.path.basename(input_path), "FAILED", str(e))
            return False

    def process_excel(self, input_path, output_path):
        try:
            print(f"[INFO] Restoring EXCEL: {os.path.basename(input_path)} ...")
            wb = openpyxl.load_workbook(input_path)

            if self.regex_pattern:
                for ws in wb:
                    count = 0
                    for row in ws.iter_rows():
                        for cell in row:
                            if isinstance(cell.value, str) and cell.value.strip():
                                cell.value = self.regex_pattern.sub(self._replace_func, cell.value)
                        count += 1
                        if count % 2000 == 0:
                            print(f"       Processed {count} rows in sheet: {ws.title}")

            out_dir = os.path.dirname(output_path)
            if not os.path.exists(out_dir): os.makedirs(out_dir)
            wb.save(output_path)
            print(f"[SUCCESS] Saved: {os.path.basename(output_path)}")
            self.log_operation(os.path.basename(input_path), "SUCCESS", "EXCEL Restoration successful")
            return True
        except Exception as e:
            print(f"[ERROR] EXCEL restoration failed: {e}")
            self.log_operation(os.path.basename(input_path), "FAILED", str(e))
            return False


if __name__ == "__main__":
    app = Restorer(DB_PATH, MASTER_KEY)
    # The run logic is typically triggered via GUI, testing scope omitted here