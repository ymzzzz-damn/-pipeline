import pandas as pd
from docx import Document
import os
import time
import re

# 配置
CSV_PATH = "./input_life/emails.csv"
OUTPUT_DIR = "./input_life"
EMAILS_PER_DOCX = 5000


def remove_control_characters(text):
    """
    清洗函数：移除 XML 不兼容的控制字符
    保留：\x09 (Tab), \x0A (换行), \x0D (回车)
    移除：\x00-\x08, \x0B-\x0C, \x0E-\x1F
    """
    if not isinstance(text, str):
        return ""
    # 使用正则替换掉非法的 ASCII 控制字符
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)


def convert_all_emails():
    if not os.path.exists(CSV_PATH):
        print(f"找不到 {CSV_PATH}，请确认文件位置。")
        return

    print(f"准备处理全量数据，策略：每 {EMAILS_PER_DOCX} 封邮件生成一个 Word 文档...")
    chunk_iterator = pd.read_csv(CSV_PATH, chunksize=EMAILS_PER_DOCX)


    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)

    for i, chunk in enumerate(chunk_iterator):
        file_index = i + 1

        # 如果你想跳过前4个已经生成好的，可以取消下面两行的注释
        # if file_index <= 4:
        #     continue

        doc = Document()
        doc.add_heading(f'Enron Email Dataset - Batch {file_index}', 0)

        print(
            f"正在生成第 {file_index} 个文件 (包含第 {i * EMAILS_PER_DOCX} - {(i + 1) * EMAILS_PER_DOCX} 条数据)...")

        for index, row in chunk.iterrows():
            content = row.get('message', '')

            # 【关键修改】先清洗，再截断
            # 1. 转为字符串
            # 2. 移除非法字符
            # 3. 截取前 5000 字
            clean_content = remove_control_characters(str(content))[:5000]

            if clean_content.strip():
                try:
                    doc.add_paragraph(f"--- Email ID: {index} ---")
                    doc.add_paragraph(clean_content)
                    doc.add_paragraph("\n")
                except Exception as e:
                    print(f" 跳过损坏邮件 ID {index}: {e}")

        save_name = f"Enron_Batch_{file_index:03d}.docx"
        save_path = os.path.join(OUTPUT_DIR, save_name)
        doc.save(save_path)

        print(f" 已保存: {save_name}")

    print(f"\n转换完成！")


if __name__ == "__main__":
    convert_all_emails()